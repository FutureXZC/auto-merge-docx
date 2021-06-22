import os
from docx import Document
from docxcompose.composer import Composer


def merge_doc(source_file_path_list, target_file_path):
    page_break_doc = Document()
    page_break_doc.add_page_break()
    target_doc = Document(source_file_path_list[0])
    target_composer = Composer(target_doc)
    for i in range(len(source_file_path_list)):
        #跳过第一个作为模板的文件
        if i == 0:
            continue
        #填充分页符文档
        target_composer.append(page_break_doc)
        #拼接文档内容
        f = source_file_path_list[i]
        target_composer.append(Document(f))
        #保存目标文档
        target_composer.save(target_file_path)


if __name__ == '__main__':
    source_path = os.getcwd() + '\\\\files\\\\'
    target_file = os.getcwd() + '\\\\merge_result.docx'
    source_file_list = os.listdir(source_path)
    source_file_list_all = []
    for file in source_file_list:
        source_file_list_all.append(source_path + file)
    merge_doc(source_file_list_all, target_file)
